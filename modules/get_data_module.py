import io
import zipfile
from typing import List
from PyPDF2 import PdfReader


def get_data(files: List[bytes], urls: List[str]) -> str:
    extracted_text = []
    if files:
        # Process each file blob
        for blob in files:
            if blob.startswith(b'%PDF-'):
                # Handle PDF
                try:
                    pdf_file = io.BytesIO(blob)
                    pdf = PdfReader(pdf_file)
                    text = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(page_text)
                    extracted_text.append('\n'.join(text))
                except Exception as e:
                    print(f"Error processing PDF: {e}")
            elif len(blob) >= 4 and blob[:4] == b'PK\x03\x04':
                # Handle ZIP-based formats (DOCX or XLSX)
                try:
                    with zipfile.ZipFile(io.BytesIO(blob)) as zip_file:
                        if 'word/document.xml' in zip_file.namelist():
                            # Handle DOCX
                            try:
                                from docx import Document
                                doc = Document(io.BytesIO(blob))
                                text = []
                                for para in doc.paragraphs:
                                    text.append(para.text)
                                for table in doc.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            text.append(cell.text)
                                extracted_text.append('\n'.join(text))
                            except Exception as e:
                                print(f"Error processing DOCX: {e}")
                        elif 'xl/workbook.xml' in zip_file.namelist():
                            # Handle XLSX
                            try:
                                from openpyxl import load_workbook
                                xlsx_file = io.BytesIO(blob)
                                wb = load_workbook(xlsx_file, data_only=True)
                                text = []
                                for sheet in wb:
                                    for row in sheet.iter_rows(values_only=True):
                                        row_text = [str(cell) for cell in row if cell is not None]
                                        text.append(' '.join(row_text))
                                extracted_text.append('\n'.join(text))
                            except Exception as e:
                                print(f"Error processing XLSX: {e}")
                except Exception as e:
                    print(f"Error processing ZIP file: {e}")
            else:
                print("Unsupported file format")

    # Process each URL
    for url in urls:
        try:
            import requests
            from bs4 import BeautifulSoup

            response = requests.get(url, timeout=10)
            response.raise_for_status()

            content_type = response.headers.get('Content-Type', '')
            if 'text/html' in content_type:
                soup = BeautifulSoup(response.content, 'html.parser')
                text = soup.get_text(separator='\n', strip=True)
                extracted_text.append(text)
            else:
                text = response.text
                extracted_text.append(text)
        except Exception as e:
            print(f"Error processing URL {url}: {e}")

    return '\n\n'.join(extracted_text)


if __name__ == '__main__':
    from util import load_source_cache, save_source, text_to_chunks
    from langchain_community.embeddings import OllamaEmbeddings
    from langchain_community.vectorstores import FAISS
    from dotenv import load_dotenv
    import json
    import os
    from pathlib import Path

    load_dotenv()
    emb_config_list = json.loads(os.getenv("EMB_CONFIG_LIST"))

    text=get_data([], urls=["http://tendium.ai/en/"])
    text_documents = text_to_chunks(text)
    faiss_vectorstore = FAISS.from_documents(text_documents,
                                             OllamaEmbeddings(base_url=emb_config_list[0]["base_url"],
                                                              model=emb_config_list[0]["model"]))
    save_path = Path("../data")
    save_source(faiss_vectorstore, text_documents, save_path, "tendium")